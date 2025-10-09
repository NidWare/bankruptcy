from flask import Flask, render_template, request, send_file, flash, redirect, url_for
from docx import Document
from datetime import datetime
import io
import os
import copy
import zipfile

app = Flask(__name__)
app.secret_key = 'your-secret-key-here'  # В продакшене используйте настоящий секретный ключ

def replace_in_runs_preserve_formatting(paragraph, replacements):
    """
    Заменяет плейсхолдеры с максимальным сохранением форматирования
    """
    if not paragraph.runs:
        return
    
    # Собираем информацию о каждом символе и его форматировании
    char_data = []
    
    for run in paragraph.runs:
        for char in run.text:
            char_data.append({
                'char': char,
                'run': run,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'bold': run.font.bold,
                'italic': run.font.italic,
                'underline': run.font.underline,
                'color': run.font.color.rgb if run.font.color.rgb else None
            })
    
    if not char_data:
        return
    
    # Получаем исходный текст
    original_text = "".join(cd['char'] for cd in char_data)
    
    # Выполняем замены
    new_text = original_text
    for placeholder, replacement in replacements.items():
        new_text = new_text.replace(placeholder, replacement)
    
    if new_text == original_text:
        return  # Нет изменений
    
    # Удаляем все существующие runs
    for run in paragraph.runs[:]:
        paragraph._element.remove(run._element)
    
    # Создаем новые runs с сохранением форматирования
    if not new_text:
        return
    
    # Берем форматирование первого символа как базовое
    base_formatting = char_data[0] if char_data else None
    
    # Создаем новый run с базовым форматированием
    new_run = paragraph.add_run(new_text)
    
    if base_formatting:
        if base_formatting['font_name']:
            new_run.font.name = base_formatting['font_name']
        if base_formatting['font_size']:
            new_run.font.size = base_formatting['font_size']
        if base_formatting['bold'] is not None:
            new_run.font.bold = base_formatting['bold']
        if base_formatting['italic'] is not None:
            new_run.font.italic = base_formatting['italic']
        if base_formatting['underline'] is not None:
            new_run.font.underline = base_formatting['underline']
        if base_formatting['color']:
            new_run.font.color.rgb = base_formatting['color']


def replace_placeholders_advanced(doc, replacements):
    """
    Улучшенная замена плейсхолдеров с обработкой всех элементов документа
    """
    # Обрабатываем все параграфы
    for paragraph in doc.paragraphs:
        replace_in_runs_preserve_formatting(paragraph, replacements)
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Обрабатываем каждый параграф в ячейке
                for paragraph in cell.paragraphs:
                    replace_in_runs_preserve_formatting(paragraph, replacements)
    
    # Обрабатываем заголовки и колонтитулы
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                replace_in_runs_preserve_formatting(paragraph, replacements)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                replace_in_runs_preserve_formatting(paragraph, replacements)


def add_creditors_rows_improved(doc, creditors):
    """
    Улучшенное добавление строк кредиторов с сохранением форматирования
    """
    for table_idx, table in enumerate(doc.tables):
        # Ищем таблицу с кредиторами (проверяем заголовки)
        if len(table.rows) > 0:
            header_row = table.rows[0]
            header_text = " ".join(cell.text.strip().lower() for cell in header_row.cells)
            
            # Если это таблица кредиторов (первая найденная)
            if "кредитор" in header_text and ("обязательство" in header_text or "денежным обязательствам" in header_text):
                print(f"Найдена таблица кредиторов с {len(table.rows)} строками")
                
                # Находим существующие строки с данными кредиторов (строки 1.1, 1.2 и т.д.)
                creditor_rows = {}  # {номер_строки: индекс_кредитора}
                
                for i, row in enumerate(table.rows):
                    row_text = " ".join(cell.text.strip() for cell in row.cells)
                    
                    # Ищем строки с номерацией 1.1, 1.2 и т.д. (денежные обязательства)
                    if row_text.startswith("1."):
                        try:
                            # Извлекаем номер кредитора (1.1 -> 1, 1.2 -> 2)
                            creditor_num = int(row_text.split()[0].split('.')[1])
                            creditor_rows[i] = creditor_num
                            print(f"Найдена строка кредитора {creditor_num} в позиции {i}")
                        except (ValueError, IndexError):
                            pass
                
                print(f"Найдено существующих строк кредиторов: {len(creditor_rows)}")
                
                # Список строк для удаления (если кредиторов меньше, чем строк в шаблоне)
                rows_to_delete = []
                
                # Заменяем существующие строки кредиторов
                for row_idx, creditor_num in creditor_rows.items():
                    if creditor_num <= len(creditors):  # Есть данные для замены
                        creditor = creditors[creditor_num - 1]  # creditor_num начинается с 1
                        cells = table.rows[row_idx].cells
                        
                        if len(cells) >= 8:
                            cells[0].text = f"1.{creditor_num}"
                            cells[1].text = creditor.get("Содержание обязательства", "")
                            cells[2].text = creditor.get("Кредитор", "")
                            cells[3].text = creditor.get("Место нахождения", "")
                            cells[4].text = creditor.get("Основание", "")
                            # Форматируем числовые значения с пробелами
                            cells[5].text = format_amount(creditor.get("Сумма обязательства", ""))
                            cells[6].text = format_amount(creditor.get("Задолженность", ""))
                            cells[7].text = format_amount(creditor.get("Штрафы", ""))
                            print(f"Заменен кредитор {creditor_num}: {creditor.get('Кредитор', 'Неизвестно')}")
                    else:
                        # Кредитора нет - нужно удалить эту строку
                        rows_to_delete.append(row_idx)
                        print(f"Строка кредитора {creditor_num} в позиции {row_idx} будет удалена (нет данных)")
                
                # Если кредиторов больше, чем существующих строк - добавляем новые
                max_existing_creditor = max(creditor_rows.values()) if creditor_rows else 0
                
                if len(creditors) > max_existing_creditor:
                    print(f"Нужно добавить еще {len(creditors) - max_existing_creditor} кредиторов")
                    
                    # Находим и временно сохраняем строки раздела "2. Обязательные платежи"
                    section_2_rows = []
                    section_2_start = None
                    
                    for i, row in enumerate(table.rows):
                        row_text = " ".join(cell.text.strip() for cell in row.cells)
                        
                        # Если нашли начало раздела 2
                        if section_2_start is None and (row_text.strip() == "2" or (row_text.strip().startswith("2") and "обязательные платежи" in row_text.lower())):
                            section_2_start = i
                            print(f"Найден раздел 2 в позиции {i}")
                        
                        # Если мы в разделе 2, сохраняем полную структуру строк
                        if section_2_start is not None and i >= section_2_start:
                            # Сохраняем полный XML элемент строки
                            section_2_rows.append(copy.deepcopy(row._element))
                    
                    # Удаляем строки раздела 2 (в обратном порядке)
                    if section_2_rows:
                        print(f"Временно удаляем {len(section_2_rows)} строк раздела 2")
                        for i in range(len(section_2_rows)):
                            # Удаляем с конца, начиная с section_2_start
                            row_to_remove = len(table.rows) - 1
                            if row_to_remove >= section_2_start:
                                table._element.remove(table.rows[row_to_remove]._element)
                    
                    # Теперь добавляем недостающих кредиторов
                    for i in range(max_existing_creditor, len(creditors)):
                        creditor = creditors[i]
                        
                        # Добавляем строку в конец таблицы (теперь без раздела 2)
                        new_row = table.add_row()
                        cells = new_row.cells
                        
                        if len(cells) >= 8:
                            cells[0].text = f"1.{i + 1}"
                            cells[1].text = creditor.get("Содержание обязательства", "")
                            cells[2].text = creditor.get("Кредитор", "")
                            cells[3].text = creditor.get("Место нахождения", "")
                            cells[4].text = creditor.get("Основание", "")
                            # Форматируем числовые значения с пробелами
                            cells[5].text = format_amount(creditor.get("Сумма обязательства", ""))
                            cells[6].text = format_amount(creditor.get("Задолженность", ""))
                            cells[7].text = format_amount(creditor.get("Штрафы", ""))
                            print(f"Добавлен кредитор {i + 1}: {creditor.get('Кредитор', 'Неизвестно')}")
                        else:
                            print(f"Недостаточно ячеек для кредитора {i + 1}: {len(cells)} < 8")
                    
                    # Восстанавливаем строки раздела 2 с сохранением оригинальной структуры
                    if section_2_rows:
                        print(f"Восстанавливаем {len(section_2_rows)} строк раздела 2")
                        table_element = table._element
                        for row_element in section_2_rows:
                            # Добавляем сохраненный элемент строки обратно в таблицу
                            table_element.append(row_element)
                
                # Удаляем лишние строки кредиторов (если кредиторов меньше, чем было в шаблоне)
                # Удаляем в обратном порядке, чтобы индексы не сбивались
                for row_idx in sorted(rows_to_delete, reverse=True):
                    print(f"Удаляем лишнюю строку кредитора в позиции {row_idx}")
                    table._element.remove(table.rows[row_idx]._element)
                
                break  # Обработали первую таблицу кредиторов


def process_document_in_memory(template_path, replacements, creditors=None):
    """
    Обрабатывает документ в памяти и возвращает байты обработанного документа
    """
    # Проверяем формат файла
    if template_path.endswith('.doc') and not template_path.endswith('.docx'):
        raise ValueError(f"Файл '{template_path}' имеет старый формат .doc. "
                        f"Пожалуйста, откройте его в Microsoft Word и сохраните как .docx формат, "
                        f"либо используйте LibreOffice для конвертации.")
    
    # Загружаем шаблон
    doc = Document(template_path)
    
    # Используем улучшенную функцию замены плейсхолдеров
    replace_placeholders_advanced(doc, replacements)
    
    # Если есть данные кредиторов и это документ со списком кредиторов, добавляем их
    if creditors and template_path.endswith('list-of-creditors.docx'):
        add_creditors_rows_improved(doc, creditors)
    
    # Сохраняем документ в память
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io


def format_amount(amount_str):
    """
    Форматирует сумму в формат с разделителями тысяч и копейками
    Пример: "1000000" -> "1 000 000,00"
    """
    try:
        # Удаляем все пробелы и запятые
        amount_str = amount_str.replace(' ', '').replace(',', '.').replace('\xa0', '')
        amount = float(amount_str)
        # Форматируем с двумя знаками после запятой
        formatted = f"{amount:,.2f}".replace(',', ' ').replace('.', ',')
        return formatted
    except (ValueError, AttributeError):
        return amount_str


def format_judge_name(full_name):
    """
    Форматирует полное ФИО судьи в формат "Фамилия И.О."
    Пример: "Иванова Мария Петровна" -> "Иванова М.П."
    """
    if not full_name:
        return ""
    
    parts = full_name.strip().split()
    if len(parts) == 0:
        return ""
    elif len(parts) == 1:
        return parts[0]  # Только фамилия
    elif len(parts) == 2:
        # Фамилия Имя
        return f"{parts[0]} {parts[1][0]}."
    else:
        # Фамилия Имя Отчество
        return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."


def calculate_total_debt(creditors):
    """
    Вычисляет общую сумму задолженности из всех кредиторов
    """
    total = 0.0
    for creditor in creditors:
        try:
            debt = creditor.get('Задолженность', '0').replace(' ', '').replace(',', '.').replace('\xa0', '')
            total += float(debt)
        except (ValueError, AttributeError):
            pass
    return total


def generate_initial_documents_archive(replacements, creditors, surname, name, current_date):
    """
    Генерирует первоначальные документы (БЕЗ номера дела):
    - заявление о банкротстве
    - список кредиторов
    - опись имущества
    """
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # Генерируем заявление о банкротстве
        bankruptcy_template = "zayav.docx"
        if os.path.exists(bankruptcy_template):
            try:
                bankruptcy_doc = process_document_in_memory(bankruptcy_template, replacements)
                bankruptcy_filename = f"bankruptcy_application_{surname}_{name}_{current_date.strftime('%Y%m%d')}.docx"
                zip_file.writestr(bankruptcy_filename, bankruptcy_doc.getvalue())
                print(f"✅ Создано заявление о банкротстве: {bankruptcy_filename}")
            except Exception as e:
                print(f"❌ Ошибка при создании заявления о банкротстве: {e}")
        else:
            print("⚠️ Шаблон заявления о банкротстве (zayav.docx) не найден")
        
        # Генерируем список кредиторов
        creditors_template = "list-of-creditors.docx"
        if os.path.exists(creditors_template):
            try:
                creditors_doc = process_document_in_memory(creditors_template, replacements, creditors)
                creditors_filename = f"list_of_creditors_{surname}_{name}_{current_date.strftime('%Y%m%d')}.docx"
                zip_file.writestr(creditors_filename, creditors_doc.getvalue())
                print(f"✅ Создан список кредиторов: {creditors_filename}")
            except Exception as e:
                print(f"❌ Ошибка при создании списка кредиторов: {e}")
        else:
            print("⚠️ Шаблон списка кредиторов (list-of-creditors.docx) не найден")
        
        # Генерируем опись имущества
        properties_template = "properties.docx"
        if os.path.exists(properties_template):
            try:
                properties_doc = process_document_in_memory(properties_template, replacements)
                properties_filename = f"properties_{surname}_{name}_{current_date.strftime('%Y%m%d')}.docx"
                zip_file.writestr(properties_filename, properties_doc.getvalue())
                print(f"✅ Создана опись имущества: {properties_filename}")
            except Exception as e:
                print(f"❌ Ошибка при создании описи имущества: {e}")
        else:
            print("⚠️ Шаблон описи имущества (properties.docx) не найден")
    
    zip_buffer.seek(0)
    return zip_buffer


def generate_case_documents_archive(replacements, surname, name, current_date):
    """
    Генерирует документы после открытия дела (С номером дела):
    - информационное сообщение
    - заявление от СРО
    - заявление о согласии арбитражного управляющего
    
    Примечание: Кредиторы не требуются для этих документов
    """
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # Генерируем информационное сообщение
        inform_template = "inform-message.docx"
        if os.path.exists(inform_template):
            try:
                inform_doc = process_document_in_memory(inform_template, replacements)
                inform_filename = f"inform_message_{surname}_{name}_{current_date.strftime('%Y%m%d')}.docx"
                zip_file.writestr(inform_filename, inform_doc.getvalue())
                print(f"✅ Создано информационное сообщение: {inform_filename}")
            except Exception as e:
                print(f"❌ Ошибка при создании информационного сообщения: {e}")
        else:
            print("⚠️ Шаблон информационного сообщения (inform-message.docx) не найден")
        
        # Генерируем заявление от СРО
        # Используем обновленную версию zayavSRO1.docx
        sro_template = "zayavSRO1.docx"
        
        if os.path.exists(sro_template):
            try:
                sro_doc = process_document_in_memory(sro_template, replacements)
                sro_filename = f"sro_application_{surname}_{name}_{current_date.strftime('%Y%m%d')}.docx"
                zip_file.writestr(sro_filename, sro_doc.getvalue())
                print(f"✅ Создано заявление от СРО: {sro_filename}")
            except Exception as e:
                print(f"❌ Ошибка при создании заявления от СРО: {e}")
        else:
            print("⚠️ Шаблон заявления от СРО (zayavSRO1.docx) не найден")
        
        # Генерируем заявление о согласии арбитражного управляющего
        agreement_template = "zayavAgreement.docx"
        if os.path.exists(agreement_template):
            try:
                agreement_doc = process_document_in_memory(agreement_template, replacements)
                agreement_filename = f"agreement_{surname}_{name}_{current_date.strftime('%Y%m%d')}.docx"
                zip_file.writestr(agreement_filename, agreement_doc.getvalue())
                print(f"✅ Создано заявление о согласии: {agreement_filename}")
            except Exception as e:
                print(f"❌ Ошибка при создании заявления о согласии: {e}")
        else:
            print("⚠️ Шаблон заявления о согласии (zayavAgreement.docx) не найден")
    
    zip_buffer.seek(0)
    return zip_buffer


@app.route('/')
def home():
    """Главная страница с выбором типа документов"""
    return render_template('home.html')


@app.route('/initial', methods=['GET', 'POST'])
def initial_documents():
    """Форма и генерация первоначальных документов (без номера дела)"""
    if request.method == 'POST':
        # Получаем основные данные из формы
        surname = request.form.get('surname', '').strip()
        name = request.form.get('name', '').strip()
        patronymic = request.form.get('patronymic', '').strip()
        surname_genitive = request.form.get('surname_genitive', '').strip()  # Родительный падеж
        name_genitive = request.form.get('name_genitive', '').strip()
        patronymic_genitive = request.form.get('patronymic_genitive', '').strip()
        surname_dative = request.form.get('surname_dative', '').strip()  # Дательный падеж
        name_dative = request.form.get('name_dative', '').strip()
        patronymic_dative = request.form.get('patronymic_dative', '').strip()
        birth_date = request.form.get('birth_date', '').strip()
        birth_place = request.form.get('birth_place', '').strip()
        passport_series = request.form.get('passport_series', '').strip()
        passport_number = request.form.get('passport_number', '').strip()
        passport_issued_by = request.form.get('passport_issued_by', '').strip()
        passport_issue_date = request.form.get('passport_issue_date', '').strip()
        inn = request.form.get('inn', '').strip()
        snils = request.form.get('snils', '').strip()
        
        # Получаем данные адреса
        region = request.form.get('region', '').strip()
        district = request.form.get('district', '').strip()
        city = request.form.get('city', '').strip()
        street = request.form.get('street', '').strip()
        house_number = request.form.get('house_number', '').strip()
        building_number = request.form.get('building_number', '').strip()
        apartment_number = request.form.get('apartment_number', '').strip()
        registered_address = request.form.get('registered_address', '').strip()
        
        # Получаем данные о долге
        debt_amount_digits = request.form.get('debt_amount_digits', '').strip()
        debt_amount_words = request.form.get('debt_amount_words', '').strip()
        
        # Получаем данные о госпошлине (по умолчанию 300 рублей для банкротства физлиц)
        state_duty = request.form.get('state_duty', '300').strip()
        if not state_duty:
            state_duty = '300'
        
        # Получаем данные о деле и судье (для документа СРО)
        case_number = request.form.get('case_number', '').strip()
        judge_name = request.form.get('judge_name', '').strip()
        
        # Валидация основных полей
        required_fields = [surname, name, patronymic, surname_genitive, name_genitive, 
                          patronymic_genitive, surname_dative, name_dative, patronymic_dative,
                          birth_date, birth_place, passport_series, passport_number,
                          passport_issued_by, passport_issue_date,
                          inn, snils, region, street, house_number, registered_address,
                          debt_amount_digits, debt_amount_words]
        if not all(required_fields):
            flash('Все основные поля обязательны для заполнения', 'error')
            return redirect(url_for('initial_documents'))
        
        if len(inn) != 12 or not inn.isdigit():
            flash('ИНН должен содержать ровно 12 цифр', 'error')
            return redirect(url_for('initial_documents'))
        
        # Получаем данные о кредиторах с детальной информацией
        creditors = []
        creditor_num = 1
        while True:
            creditor_name = request.form.get(f'creditor_name_{creditor_num}', '').strip()
            creditor_address = request.form.get(f'creditor_address_{creditor_num}', '').strip()
            obligation_content = request.form.get(f'obligation_content_{creditor_num}', '').strip()
            obligation_basis = request.form.get(f'obligation_basis_{creditor_num}', '').strip()
            obligation_amount = request.form.get(f'obligation_amount_{creditor_num}', '').strip()
            debt_amount = request.form.get(f'debt_amount_{creditor_num}', '').strip()
            penalties = request.form.get(f'penalties_{creditor_num}', '').strip()
            
            # Если основные поля кредитора пусты, прекращаем поиск
            if not creditor_name and not creditor_address:
                break
            
            # Проверяем, что все поля кредитора заполнены
            creditor_fields = [creditor_name, creditor_address, obligation_content, 
                              obligation_basis, obligation_amount, debt_amount, penalties]
            if not all(creditor_fields):
                flash(f'Заполните все поля для кредитора {creditor_num}', 'error')
                return redirect(url_for('initial_documents'))
            
            creditors.append({
                'name': creditor_name,
                'address': creditor_address,
                'Содержание обязательства': obligation_content,
                'Кредитор': creditor_name,
                'Место нахождения': creditor_address,
                'Основание': obligation_basis,
                'Сумма обязательства': obligation_amount,
                'Задолженность': debt_amount,
                'Штрафы': penalties
            })
            creditor_num += 1
        
        if not creditors:
            flash('Необходимо указать хотя бы одного кредитора', 'error')
            return redirect(url_for('initial_documents'))
        
        try:
            # Получаем текущую дату
            current_date = datetime.now()
            
            # Получаем первые буквы имени и отчества
            first_letter_name = name[0].upper() if name else ''
            first_letter_patronymic = patronymic[0].upper() if patronymic else ''
            
            # Получаем название месяца на русском языке
            months_ru = [
                '', 'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
                'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря'
            ]
            month_name = months_ru[current_date.month]
            
            # Форматируем дату рождения
            try:
                birth_date_obj = datetime.strptime(birth_date, '%Y-%m-%d')
                formatted_birth_date = birth_date_obj.strftime('%d.%m.%Y')
            except ValueError:
                formatted_birth_date = birth_date
            
            # Форматируем дату выдачи паспорта
            try:
                passport_issue_date_obj = datetime.strptime(passport_issue_date, '%Y-%m-%d')
                formatted_passport_issue_date = passport_issue_date_obj.strftime('%d.%m.%Y')
            except ValueError:
                formatted_passport_issue_date = passport_issue_date
            
            # Формируем полное ФИО должника в разных падежах
            full_name = f"{surname} {name} {patronymic}"
            full_name_genitive = f"{surname_genitive} {name_genitive} {patronymic_genitive}"
            full_name_dative = f"{surname_dative} {name_dative} {patronymic_dative}"
            
            # Формируем паспортные данные
            passport_full = f"{passport_series} {passport_number}"
            
            # Формируем полную строку с датой рождения и местом рождения
            birth_info = f"{formatted_birth_date} г.р., место рождения: {birth_place}"
            
            # Формируем полную строку с паспортными данными
            passport_info = f"Паспорт РФ {passport_full}\nВыдан {passport_issued_by}\nдата выдачи: {formatted_passport_issue_date}"
            
            # Вычисляем общую сумму задолженности
            total_debt = calculate_total_debt(creditors)
            formatted_total_debt = format_amount(str(total_debt))
            
            # Подготавливаем основные замены (включая новые поля из list-of-creditors-final.py)
            replacements = {
                # Основные персональные данные (именительный падеж)
                "{Фамилия}": surname,
                "{Фамилия} ": surname + " ",
                "{фамилия}": surname,
                "{фамилия} ": surname + " ",
                "{Имя}": name,
                "{имя}": name,
                "{Отчество}": patronymic,
                "{отчество}": patronymic,
                "{ФИО}": full_name,
                "{ФИО должника}": full_name,
                
                # ФИО в родительном падеже (кого? чего?)
                "{Фамилия родительный}": surname_genitive,
                "{Имя родительный}": name_genitive,
                "{Отчество родительный}": patronymic_genitive,
                "{ФИО родительный}": full_name_genitive,
                
                # ФИО в дательном падеже (кому? чему?)
                "{Фамилия дательный}": surname_dative,
                "{Имя дательный}": name_dative,
                "{Отчество дательный}": patronymic_dative,
                "{ФИО дательный}": full_name_dative,
                
                # Дата и место рождения
                "{дата рождения}": formatted_birth_date,
                "{dd.mm.yyyy дата рождения}": formatted_birth_date,
                "{место рождения}": birth_place,
                "{дата и место рождения}": birth_info,
                
                # Паспортные данные
                "{паспортные данные}": passport_info,
                "{серия и номер паспорта}": passport_full,
                "{паспорт серия}": passport_series,
                "{паспорт номер}": passport_number,
                "{паспорт кем выдан}": passport_issued_by,
                "{паспорт дата выдачи}": formatted_passport_issue_date,
                
                # ИНН и СНИЛС
                "{ИНН}": inn,
                "{СНИЛС}": snils,
                
                # Детализированный адрес
                "{субъект РФ}": region,
                "{район (при наличии)}": district,
                "{город (при наличии)}": city,
                "{населенный пункт}": city,
                "{улица}": street,
                "{номер дома}": house_number,
                "{номер корпуса}": building_number,
                "{номер корпуса (может быть пустым)}": building_number,
                "{номер квартиры}": apartment_number,
                "{номер квартиры может быть пустым}": apartment_number,
                "{Зарегистрирован по адресу}": registered_address,
                
                # Финансовая информация
                "{сумма долга цифрами}": debt_amount_digits,
                "{сумма долга буквами}": debt_amount_words,
                "{общая сумма задолженности}": f"{formatted_total_debt} рублей",
                "{сумма требований}": formatted_total_debt,
                "{госпошлина}": state_duty,
                
                # Дата и подпись
                "{dd}.{mm}.{yyyy} г.": f"{current_date.day:02d}.{current_date.month:02d}.{current_date.year} г.",
                "{dd}.{mm}.{yyyy}г.": f"{current_date.day:02d}.{current_date.month:02d}.{current_date.year}г.",
                "{dd}": f"{current_date.day:02d}",
                "{mm}": f"{current_date.month:02d}",
                "{month name}": month_name,
                "{yyyy}": str(current_date.year),
                "{дата}": current_date.strftime("%d.%m.%Y"),
                "{число месяца}": str(current_date.day),
                "{месяц}": month_name,
                "{год}": str(current_date.year),
                "{Первая буква имени}": first_letter_name,
                "{первая буква отчества}": first_letter_patronymic,
                "{Фамилия и первые буквы имени и отчества}": f"{surname} {first_letter_name}.{first_letter_patronymic}.",
                
                # Плейсхолдеры для информационного сообщения
                "{ИНН ДОЛЖНИКА}": inn,
                "{СНИЛС ДОЛЖНИКА}": snils,
                "{номер дела}": case_number if case_number else "",
                "{месторасположение должника}": registered_address,
                "{сумма требований к должнику}": f"{formatted_total_debt} рублей",
                
                # Плейсхолдеры для заявления от СРО
                "{дело}": case_number if case_number else "",
                "{судья}": judge_name if judge_name else "",
                
                # Кредиторы (общие плейсхолдеры)
                "{Наименование кредитора}": creditors[0]['name'] if creditors else "",
                "{Почтовый индекс и адрес}": creditors[0]['address'] if creditors else "",
                "{место нахождения кредитора}": creditors[0]['address'] if creditors else "",
                "{основание возникновения}": creditors[0]['Основание'] if creditors else "",
                "{сумма обязательства}": creditors[0]['Сумма обязательства'] if creditors else "",
                "{сумма задолженности}": creditors[0]['Задолженность'] if creditors else "",
                "{штрафы + пени}": creditors[0]['Штрафы'] if creditors else "",
            }
            
            # Добавляем замены для первых кредиторов (из оригинального кода)
            if len(creditors) >= 1:
                replacements.update({
                    "{кредит1}": creditors[0]['Содержание обязательства'],
                    "{кредитор1}": creditors[0]['Кредитор'],
                })
            
            if len(creditors) >= 2:
                replacements.update({
                    "{кредит2}": creditors[1]['Содержание обязательства'],
                    "{кредитор2}": creditors[1]['Кредитор'],
                })
            
            # Добавляем замены для всех кредиторов с номерами
            for i, creditor in enumerate(creditors, 1):
                replacements[f"{{Кредитор {i}}}"] = f"Кредитор {i}"
                replacements[f"{{Кредитор n}}"] = f"Кредитор {i}" if i == 1 else replacements.get(f"{{Кредитор n}}", f"Кредитор {i}")
                if i > 1:
                    replacements[f"{{Кредитор n+{i-1}}}"] = f"Кредитор {i}"
            
            # Проверяем наличие хотя бы одного шаблона
            bankruptcy_template = "zayav.docx"
            creditors_template = "list-of-creditors.docx"
            properties_template = "properties.docx"
            
            if not os.path.exists(bankruptcy_template) and not os.path.exists(creditors_template):
                flash('Файлы шаблонов не найдены', 'error')
                return redirect(url_for('initial_documents'))
            
            # Создаем ZIP-архив с первоначальными документами
            zip_buffer = generate_initial_documents_archive(replacements, creditors, surname, name, current_date)
            filename = f"initial_documents_{surname}_{name}_{current_date.strftime('%Y%m%d')}.zip"
            
            return send_file(
                zip_buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='application/zip'
            )
            
        except Exception as e:
            flash(f'Ошибка при обработке документа: {str(e)}', 'error')
            return redirect(url_for('initial_documents'))
    
    return render_template('initial.html')


@app.route('/with-case', methods=['GET', 'POST'])
def case_documents():
    """Форма и генерация документов после открытия дела (с номером дела)"""
    if request.method == 'POST':
        # Получаем основные данные из формы
        surname = request.form.get('surname', '').strip()
        name = request.form.get('name', '').strip()
        patronymic = request.form.get('patronymic', '').strip()
        surname_genitive = request.form.get('surname_genitive', '').strip()
        name_genitive = request.form.get('name_genitive', '').strip()
        patronymic_genitive = request.form.get('patronymic_genitive', '').strip()
        surname_dative = request.form.get('surname_dative', '').strip()
        name_dative = request.form.get('name_dative', '').strip()
        patronymic_dative = request.form.get('patronymic_dative', '').strip()
        birth_date = request.form.get('birth_date', '').strip()
        birth_place = request.form.get('birth_place', '').strip()
        passport_series = request.form.get('passport_series', '').strip()
        passport_number = request.form.get('passport_number', '').strip()
        passport_issued_by = request.form.get('passport_issued_by', '').strip()
        passport_issue_date = request.form.get('passport_issue_date', '').strip()
        inn = request.form.get('inn', '').strip()
        snils = request.form.get('snils', '').strip()
        
        # Получаем адрес (для документов после открытия дела достаточно полного адреса)
        registered_address = request.form.get('registered_address', '').strip()
        
        # Получаем общую сумму требований
        total_debt = request.form.get('total_debt', '').strip()
        
        # Получаем данные о деле и судье (ОБЯЗАТЕЛЬНЫ для этой формы)
        case_number = request.form.get('case_number', '').strip()
        judge_name = request.form.get('judge_name', '').strip()
        
        # Валидация основных полей + обязательные номер дела и судья
        required_fields = [surname, name, patronymic, surname_genitive, name_genitive,
                          patronymic_genitive, surname_dative, name_dative, patronymic_dative,
                          birth_date, birth_place, passport_series, passport_number,
                          passport_issued_by, passport_issue_date,
                          inn, snils, registered_address, total_debt,
                          case_number, judge_name]
        if not all(required_fields):
            flash('Все поля обязательны для заполнения (включая номер дела и судью)', 'error')
            return redirect(url_for('case_documents'))
        
        if len(inn) != 12 or not inn.isdigit():
            flash('ИНН должен содержать ровно 12 цифр', 'error')
            return redirect(url_for('case_documents'))
        
        try:
            current_date = datetime.now()
            
            first_letter_name = name[0].upper() if name else ''
            first_letter_patronymic = patronymic[0].upper() if patronymic else ''
            
            months_ru = [
                '', 'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
                'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря'
            ]
            month_name = months_ru[current_date.month]
            
            # Форматируем дату рождения
            try:
                birth_date_obj = datetime.strptime(birth_date, '%Y-%m-%d')
                formatted_birth_date = birth_date_obj.strftime('%d.%m.%Y')
            except ValueError:
                formatted_birth_date = birth_date
            
            # Форматируем дату выдачи паспорта
            try:
                passport_issue_date_obj = datetime.strptime(passport_issue_date, '%Y-%m-%d')
                formatted_passport_issue_date = passport_issue_date_obj.strftime('%d.%m.%Y')
            except ValueError:
                formatted_passport_issue_date = passport_issue_date
            
            # Формируем полное ФИО должника в разных падежах
            full_name = f"{surname} {name} {patronymic}"
            full_name_genitive = f"{surname_genitive} {name_genitive} {patronymic_genitive}"
            full_name_dative = f"{surname_dative} {name_dative} {patronymic_dative}"
            
            # Формируем паспортные данные
            passport_full = f"{passport_series} {passport_number}"
            
            # Формируем полную строку с датой рождения и местом рождения
            birth_info = f"{formatted_birth_date} г.р., место рождения: {birth_place}"
            
            # Формируем полную строку с паспортными данными
            passport_info = f"Паспорт РФ {passport_full}\nВыдан {passport_issued_by}\nдата выдачи: {formatted_passport_issue_date}"
            
            # Форматируем сумму требований
            formatted_total_debt = format_amount(total_debt)
            
            # Форматируем имя судьи
            judge_formatted = format_judge_name(judge_name)
            
            # Подготавливаем замены (только реально используемые в документах after case)
            replacements = {
                # Основные персональные данные
                "{ИНН}": inn,
                "{ФИО}": full_name,
                "{ФИО должника}": full_name,
                
                # ИНН и СНИЛС для информационного сообщения
                "{ИНН ДОЛЖНИКА}": inn,
                "{СНИЛС ДОЛЖНИКА}": snils,
                
                # Адрес (только полный адрес одной строкой)
                "{месторасположение должника}": registered_address,
                
                # Финансовая информация
                "{сумма требований}": formatted_total_debt,
                "{сумма требований к должнику}": f"{formatted_total_debt} рублей",
                
                # Дата
                "{dd}": f"{current_date.day:02d}",
                "{mm}": f"{current_date.month:02d}",
                "{yyyy}": str(current_date.year),
                
                # Информация о деле
                "{номер дела}": case_number,
                "{дело}": case_number,
                "{судья}": judge_name,
                "{Фамилия судьи + ее инициалы}": judge_formatted,
            }
            
            # Создаем ZIP-архив с документами после открытия дела
            zip_buffer = generate_case_documents_archive(replacements, surname, name, current_date)
            filename = f"case_documents_{surname}_{name}_{current_date.strftime('%Y%m%d')}.zip"
            
            return send_file(
                zip_buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='application/zip'
            )
            
        except Exception as e:
            flash(f'Ошибка при обработке документа: {str(e)}', 'error')
            return redirect(url_for('case_documents'))
    
    return render_template('with_case.html')


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=8080) 