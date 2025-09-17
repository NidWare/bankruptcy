from docx import Document
from datetime import datetime

def replace_in_runs(paragraph, replacements):
    """
    Заменяет плейсхолдеры в параграфе, сохраняя форматирование каждого run
    """
    # Сначала проверяем, есть ли что заменять
    full_text = "".join(run.text for run in paragraph.runs)
    has_replacements = any(key in full_text for key in replacements.keys())
    
    if not has_replacements:
        return
    
    # Создаем список кусочков текста с их форматированием
    text_pieces = []
    for run in paragraph.runs:
        if run.text:
            text_pieces.append({
                'text': run.text,
                'run': run
            })
    
    # Объединяем весь текст для поиска замен
    combined_text = "".join(piece['text'] for piece in text_pieces)
    
    # Выполняем все замены
    replaced_text = combined_text
    for key, value in replacements.items():
        replaced_text = replaced_text.replace(key, value)
    
    # Если текст изменился, нужно перераспределить его по runs
    if replaced_text != combined_text:
        # Очищаем все runs
        for run in paragraph.runs:
            run.text = ""
        
        # Если есть runs, помещаем весь новый текст в первый run
        # сохраняя его форматирование
        if paragraph.runs:
            paragraph.runs[0].text = replaced_text


def replace_in_runs_advanced(paragraph, replacements):
    """
    Продвинутая версия замены с сохранением форматирования
    """
    # Собираем информацию о всех runs и их позициях
    runs_info = []
    current_pos = 0
    
    for run in paragraph.runs:
        if run.text:
            runs_info.append({
                'run': run,
                'start': current_pos,
                'end': current_pos + len(run.text),
                'text': run.text
            })
            current_pos += len(run.text)
    
    if not runs_info:
        return
    
    # Получаем полный текст
    full_text = "".join(info['text'] for info in runs_info)
    
    # Проверяем, нужны ли замены
    replaced_text = full_text
    for key, value in replacements.items():
        replaced_text = replaced_text.replace(key, value)
    
    if replaced_text == full_text:
        return  # Нет изменений
    
    # Если текст изменился, используем простой подход:
    # сохраняем форматирование первого run и помещаем туда весь текст
    for run in paragraph.runs:
        run.text = ""
    
    if paragraph.runs:
        paragraph.runs[0].text = replaced_text


def replace_in_runs_smart(paragraph, replacements):
    """
    Умная замена с максимальным сохранением форматирования
    """
    if not paragraph.runs:
        return
    
    # Собираем информацию о runs
    runs_data = []
    for i, run in enumerate(paragraph.runs):
        runs_data.append({
            'index': i,
            'run': run,
            'original_text': run.text,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'bold': run.font.bold,
            'italic': run.font.italic,
            'underline': run.font.underline,
            'color': run.font.color.rgb if run.font.color.rgb else None
        })
    
    # Получаем полный текст
    full_text = "".join(run_data['original_text'] for run_data in runs_data)
    
    # Выполняем замены
    new_text = full_text
    for placeholder, replacement in replacements.items():
        new_text = new_text.replace(placeholder, replacement)
    
    if new_text == full_text:
        return  # Нет изменений
    
    # Находим, какие плейсхолдеры были заменены и где
    changes = []
    temp_text = full_text
    offset = 0
    
    for placeholder, replacement in replacements.items():
        while placeholder in temp_text:
            pos = temp_text.find(placeholder)
            if pos != -1:
                changes.append({
                    'start': pos + offset,
                    'end': pos + len(placeholder) + offset,
                    'old_text': placeholder,
                    'new_text': replacement,
                    'length_diff': len(replacement) - len(placeholder)
                })
                temp_text = temp_text[:pos] + replacement + temp_text[pos + len(placeholder):]
                offset += len(replacement) - len(placeholder)
            else:
                break
    
    # Если изменений слишком много или они сложные, используем простой подход
    if len(changes) > 3 or any(change['length_diff'] != 0 for change in changes):
        # Очищаем все runs кроме первого
        for i, run in enumerate(paragraph.runs):
            if i == 0:
                run.text = new_text
            else:
                run.text = ""
        return
    
    # Иначе пытаемся сохранить форматирование
    # Для простоты, если замены не меняют длину текста значительно,
    # просто заменяем текст в соответствующих runs
    char_pos = 0
    for run_data in runs_data:
        run = run_data['run']
        run_start = char_pos
        run_end = char_pos + len(run_data['original_text'])
        
        # Извлекаем соответствующую часть нового текста
        if run_end <= len(new_text):
            run.text = new_text[run_start:run_end]
        elif run_start < len(new_text):
            run.text = new_text[run_start:]
        else:
            run.text = ""
        
        char_pos = run_end
    
    # Если новый текст длиннее, добавляем остаток в последний непустой run
    if len(new_text) > char_pos:
        remaining_text = new_text[char_pos:]
        # Находим последний run с текстом
        for run in reversed(paragraph.runs):
            if run.text:
                run.text += remaining_text
                break
        else:
            # Если нет runs с текстом, добавляем в первый
            if paragraph.runs:
                paragraph.runs[0].text = new_text


def replace_in_runs_preserve_formatting(paragraph, replacements):
    """
    Заменяет плейсхолдеры с максимальным сохранением форматирования
    """
    if not paragraph.runs:
        return
    
    # Собираем информацию о каждом символе и его форматировании
    char_data = []
    
    for run in paragraph.runs:
        for char in run.text:
            char_data.append({
                'char': char,
                'run': run,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'bold': run.font.bold,
                'italic': run.font.italic,
                'underline': run.font.underline,
                'color': run.font.color.rgb if run.font.color.rgb else None
            })
    
    if not char_data:
        return
    
    # Получаем исходный текст
    original_text = "".join(cd['char'] for cd in char_data)
    
    # Выполняем замены
    new_text = original_text
    for placeholder, replacement in replacements.items():
        new_text = new_text.replace(placeholder, replacement)
    
    if new_text == original_text:
        return  # Нет изменений
    
    # Удаляем все существующие runs
    for run in paragraph.runs[:]:
        paragraph._element.remove(run._element)
    
    # Создаем новые runs с сохранением форматирования
    if new_text:
        # Находим первое место замены для определения базового форматирования
        base_formatting = None
        for placeholder in replacements.keys():
            if placeholder in original_text:
                pos = original_text.find(placeholder)
                if pos < len(char_data):
                    base_formatting = char_data[pos]
                    break
        
        if not base_formatting and char_data:
            base_formatting = char_data[0]
        
        # Создаем новый run с новым текстом
        new_run = paragraph.add_run(new_text)
        
        # Применяем форматирование
        if base_formatting:
            if base_formatting['font_name']:
                new_run.font.name = base_formatting['font_name']
            if base_formatting['font_size']:
                new_run.font.size = base_formatting['font_size']
            if base_formatting['bold'] is not None:
                new_run.font.bold = base_formatting['bold']
            if base_formatting['italic'] is not None:
                new_run.font.italic = base_formatting['italic']
            if base_formatting['underline'] is not None:
                new_run.font.underline = base_formatting['underline']
            if base_formatting['color']:
                new_run.font.color.rgb = base_formatting['color']


def replace_placeholders(doc_path, output_path, replacements):
    doc = Document(doc_path)

    # Абзацы
    for p in doc.paragraphs:
        replace_in_runs_preserve_formatting(p, replacements)

    # Таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_runs_preserve_formatting(p, replacements)

    doc.save(output_path)


if __name__ == "__main__":
    # Получаем текущую дату
    current_date = datetime.now()
    
    replacements = {
        "{Фамилия}": "Иванов",
        "{Имя}": "Алексей",
        "{Отчество}": "Петрович",
        "{ИНН}": "123456789012",
        "{dd}.{mm}.{yyyy} г.": f"{current_date.day:02d}.{current_date.month:02d}.{current_date.year} г."
    }

    replace_placeholders(
        "zayav.docx",
        "output.docx",
        replacements
    )
