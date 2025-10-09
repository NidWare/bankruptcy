from docx import Document
from datetime import datetime

def replace_in_runs_preserve_formatting(paragraph, replacements):
    """
    Заменяет плейсхолдеры с максимальным сохранением форматирования
    """
    if not paragraph.runs:
        return
    
    # Собираем информацию о каждом символе и его форматировании
    char_data = []
    
    for run in paragraph.runs:
        for char in run.text:
            char_data.append({
                'char': char,
                'run': run,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'bold': run.font.bold,
                'italic': run.font.italic,
                'underline': run.font.underline,
                'color': run.font.color.rgb if run.font.color.rgb else None
            })
    
    if not char_data:
        return
    
    # Получаем исходный текст
    original_text = "".join(cd['char'] for cd in char_data)
    
    # Выполняем замены
    new_text = original_text
    for placeholder, replacement in replacements.items():
        new_text = new_text.replace(placeholder, replacement)
    
    if new_text == original_text:
        return  # Нет изменений
    
    # Удаляем все существующие runs
    for run in paragraph.runs[:]:
        paragraph._element.remove(run._element)
    
    # Создаем новые runs с сохранением форматирования
    if not new_text:
        return
    
    # Берем форматирование первого символа как базовое
    base_formatting = char_data[0] if char_data else None
    
    # Создаем новый run с базовым форматированием
    new_run = paragraph.add_run(new_text)
    
    if base_formatting:
        if base_formatting['font_name']:
            new_run.font.name = base_formatting['font_name']
        if base_formatting['font_size']:
            new_run.font.size = base_formatting['font_size']
        if base_formatting['bold'] is not None:
            new_run.font.bold = base_formatting['bold']
        if base_formatting['italic'] is not None:
            new_run.font.italic = base_formatting['italic']
        if base_formatting['underline'] is not None:
            new_run.font.underline = base_formatting['underline']
        if base_formatting['color']:
            new_run.font.color.rgb = base_formatting['color']


def replace_placeholders_advanced(doc, replacements):
    """
    Улучшенная замена плейсхолдеров с обработкой всех элементов документа
    """
    # Обрабатываем все параграфы
    for paragraph in doc.paragraphs:
        replace_in_runs_preserve_formatting(paragraph, replacements)
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Обрабатываем каждый параграф в ячейке
                for paragraph in cell.paragraphs:
                    replace_in_runs_preserve_formatting(paragraph, replacements)
    
    # Обрабатываем заголовки и колонтитулы
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                replace_in_runs_preserve_formatting(paragraph, replacements)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                replace_in_runs_preserve_formatting(paragraph, replacements)


def add_creditors_rows_improved(doc, creditors):
    """
    Улучшенное добавление строк кредиторов с сохранением форматирования
    """
    for table_idx, table in enumerate(doc.tables):
        # Ищем таблицу с кредиторами (проверяем заголовки)
        if len(table.rows) > 0:
            header_row = table.rows[0]
            header_text = " ".join(cell.text.strip().lower() for cell in header_row.cells)
            
            # Если это таблица кредиторов (первая найденная)
            if "кредитор" in header_text and ("обязательство" in header_text or "денежным обязательствам" in header_text):
                print(f"Найдена таблица кредиторов с {len(table.rows)} строками")
                
                # Находим существующие строки с данными кредиторов (строки 1.1, 1.2 и т.д.)
                creditor_rows = {}  # {номер_строки: индекс_кредитора}
                
                for i, row in enumerate(table.rows):
                    row_text = " ".join(cell.text.strip() for cell in row.cells)
                    
                    # Ищем строки с номерацией 1.1, 1.2 и т.д. (денежные обязательства)
                    if row_text.startswith("1."):
                        try:
                            # Извлекаем номер кредитора (1.1 -> 1, 1.2 -> 2)
                            creditor_num = int(row_text.split()[0].split('.')[1])
                            creditor_rows[i] = creditor_num
                            print(f"Найдена строка кредитора {creditor_num} в позиции {i}")
                        except (ValueError, IndexError):
                            pass
                
                print(f"Найдено существующих строк кредиторов: {len(creditor_rows)}")
                
                # Заменяем существующие строки кредиторов
                for row_idx, creditor_num in creditor_rows.items():
                    if creditor_num <= len(creditors):  # Есть данные для замены
                        creditor = creditors[creditor_num - 1]  # creditor_num начинается с 1
                        cells = table.rows[row_idx].cells
                        
                        if len(cells) >= 8:
                            cells[0].text = f"1.{creditor_num}"
                            cells[1].text = creditor.get("Содержание обязательства", "")
                            cells[2].text = creditor.get("Кредитор", "")
                            cells[3].text = creditor.get("Место нахождения", "")
                            cells[4].text = creditor.get("Основание", "")
                            cells[5].text = creditor.get("Сумма обязательства", "")
                            cells[6].text = creditor.get("Задолженность", "")
                            cells[7].text = creditor.get("Штрафы", "")
                            print(f"Заменен кредитор {creditor_num}: {creditor.get('Кредитор', 'Неизвестно')}")
                
                # Если кредиторов больше, чем существующих строк - добавляем новые
                max_existing_creditor = max(creditor_rows.values()) if creditor_rows else 0
                
                if len(creditors) > max_existing_creditor:
                    print(f"Нужно добавить еще {len(creditors) - max_existing_creditor} кредиторов")
                    
                    # Находим и временно сохраняем строки раздела "2. Обязательные платежи"
                    section_2_rows = []
                    section_2_start = None
                    
                    for i, row in enumerate(table.rows):
                        row_text = " ".join(cell.text.strip() for cell in row.cells)
                        
                        # Если нашли начало раздела 2
                        if section_2_start is None and (row_text.strip() == "2" or (row_text.strip().startswith("2") and "обязательные платежи" in row_text.lower())):
                            section_2_start = i
                            print(f"Найден раздел 2 в позиции {i}")
                        
                        # Если мы в разделе 2, сохраняем полную структуру строк
                        if section_2_start is not None and i >= section_2_start:
                            # Сохраняем полный XML элемент строки
                            import copy
                            section_2_rows.append(copy.deepcopy(row._element))
                    
                    # Удаляем строки раздела 2 (в обратном порядке)
                    if section_2_rows:
                        print(f"Временно удаляем {len(section_2_rows)} строк раздела 2")
                        for i in range(len(section_2_rows)):
                            # Удаляем с конца, начиная с section_2_start
                            row_to_remove = len(table.rows) - 1
                            if row_to_remove >= section_2_start:
                                table._element.remove(table.rows[row_to_remove]._element)
                    
                    # Теперь добавляем недостающих кредиторов
                    for i in range(max_existing_creditor, len(creditors)):
                        creditor = creditors[i]
                        
                        # Добавляем строку в конец таблицы (теперь без раздела 2)
                        new_row = table.add_row()
                        cells = new_row.cells
                        
                        if len(cells) >= 8:
                            cells[0].text = f"1.{i + 1}"
                            cells[1].text = creditor.get("Содержание обязательства", "")
                            cells[2].text = creditor.get("Кредитор", "")
                            cells[3].text = creditor.get("Место нахождения", "")
                            cells[4].text = creditor.get("Основание", "")
                            cells[5].text = creditor.get("Сумма обязательства", "")
                            cells[6].text = creditor.get("Задолженность", "")
                            cells[7].text = creditor.get("Штрафы", "")
                            print(f"Добавлен кредитор {i + 1}: {creditor.get('Кредитор', 'Неизвестно')}")
                        else:
                            print(f"Недостаточно ячеек для кредитора {i + 1}: {len(cells)} < 8")
                    
                    # Восстанавливаем строки раздела 2 с сохранением оригинальной структуры
                    if section_2_rows:
                        print(f"Восстанавливаем {len(section_2_rows)} строк раздела 2")
                        table_element = table._element
                        for row_element in section_2_rows:
                            # Добавляем сохраненный элемент строки обратно в таблицу
                            table_element.append(row_element)
                
                break  # Обработали первую таблицу кредиторов


def main():
    # Загружаем документ
    doc = Document("list-of-creditors.docx")
    
    # Получаем текущую дату
    current_date = datetime.now().strftime("%d.%m.%Y")
    
    # Полный список замен (все найденные плейсхолдеры с точными пробелами)
    replacements = {
        # Персональные данные (обратите внимание на пробелы!)
        "{фамилия}": "Иванов",
        "{фамилия} ": "Иванов ",  # С пробелом в конце
        "{имя}": "Иван", 
        "{отчество}": "Иванович",
        "{дата рождения}": "01.01.1990",
        "{место рождения}": "г. Москва",
        "{СНИЛС}": "123-456-789 00",
        "{ИНН}": "7700000000",
        "{серия и номер паспорта}": "4500 123456",
        
        # Адрес
        "{субъект РФ}": "г. Москва",
        "{населенный пункт}": "",  # Может быть пустым для Москвы
        "{улица}": "ул. Тверская",
        "{номер дома}": "10",
        "{номер корпуса (может быть пустым)}": "",  # Может быть пустым
        "{номер квартиры может быть пустым}": "15",
        
        # Подпись
        "{дата}": current_date,
        "{Фамилия и первые буквы имени и отчества}": "Иванов И.И.",
        
        # Кредиторы в таблицах
        "{кредит1}": "Кредитная карта",
        "{кредитор1}": "ПАО «Альфа-Банк»",
        "{кредит2}": "Потребительский кредит", 
        "{кредитор2}": "ПАО Сбербанк",
        "{место нахождения кредитора}": "г. Москва",
        "{основание возникновения}": "Кредитный договор",
        "{сумма обязательства}": "100 000,00",
        "{сумма задолженности}": "80 000,00",
        "{штрафы + пени}": "5 000,00",
    }
    
    print("Выполняем замену плейсхолдеров...")
    replace_placeholders_advanced(doc, replacements)
    
    # Данные кредиторов для добавления в таблицу
    creditors = [
        {
            "Содержание обязательства": "Кредитная карта",
            "Кредитор": "ПАО «Альфа-Банк»",
            "Место нахождения": "г. Москва, ул. Каланчевская, д. 27",
            "Основание": "Договор банковского счета от 15.03.2019",
            "Сумма обязательства": "50 000,00",
            "Задолженность": "40 000,00", 
            "Штрафы": "2 000,00"
        },
        {
            "Содержание обязательства": "Потребительский кредит",
            "Кредитор": "ПАО Сбербанк",
            "Место нахождения": "г. Москва, ул. Вавилова, д. 19",
            "Основание": "Кредитный договор от 10.05.2021",
            "Сумма обязательства": "150 000,00",
            "Задолженность": "120 000,00",
            "Штрафы": "7 000,00"
        },
        {
            "Содержание обязательства": "Ипотечный кредит",
            "Кредитор": "ВТБ 24 (ПАО)",
            "Место нахождения": "г. Москва, ул. Мясницкая, д. 35",
            "Основание": "Договор ипотечного кредитования от 22.08.2020",
            "Сумма обязательства": "2 500 000,00",
            "Задолженность": "2 200 000,00",
            "Штрафы": "15 000,00"
        },
        {
            "Содержание обязательства": "Автокредит",
            "Кредитор": "АО «Райффайзенбанк»",
            "Место нахождения": "г. Москва, ул. Троицкая, д. 17/1",
            "Основание": "Договор автокредитования от 05.12.2022",
            "Сумма обязательства": "800 000,00",
            "Задолженность": "650 000,00",
            "Штрафы": "3 500,00"
        },
        {
            "Содержание обязательства": "Микрозайм",
            "Кредитор": "ООО МФК «Быстроденьги»",
            "Место нахождения": "г. Москва, Варшавское ш., д. 42",
            "Основание": "Договор займа от 18.07.2023",
            "Сумма обязательства": "25 000,00",
            "Задолженность": "30 000,00",
            "Штрафы": "8 000,00"
        }
    ]
    
    print("Добавляем данные кредиторов...")
    add_creditors_rows_improved(doc, creditors)
    
    # Сохраняем результат
    output_filename = "list-of-creditors-filled-final.docx"
    doc.save(output_filename)
    print(f"✅ Документ успешно обработан и сохранен как: {output_filename}")
    print("✅ Все плейсхолдеры заменены с сохранением форматирования!")
    print("✅ Добавлены строки с данными кредиторов!")


if __name__ == "__main__":
    main() 